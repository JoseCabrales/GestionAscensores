import tkinter as tk
from tkinter import messagebox
from tkinter import ttk
from tkcalendar import Calendar
import openpyxl
from docx import Document
import os
import json
from openpyxl.styles import Font

def cargar_edificios_desde_excel(archivo):
    wb = openpyxl.load_workbook(archivo)
    hoja = wb.active
    edificios_nit = {}
    edificios_nombre = []
    for fila in hoja.iter_rows(values_only=True):
        edificios_nit[fila[2]] = fila[3]
        edificios_nombre.append(fila[2])
    return edificios_nit, edificios_nombre

def actualizar_nit(event):
    edificio_seleccionado = combo_nombre_edificio.get()
    if edificio_seleccionado in edificios_nit:
        nit_edificio = edificios_nit[edificio_seleccionado]
        combo_nit_edificio.set(nit_edificio)

def agregar_ascensor():
    nit_edificio = combo_nit_edificio.get()
    if nit_edificio not in datos:
        nombre_edificio = combo_nombre_edificio.get()
        datos[nit_edificio] = {
            "nombre": nombre_edificio,
            "nit": combo_nit_edificio.get(),
            "ascensores": [],
            "estadisticas_correccion": {}  
        }

    edificio = datos[nit_edificio]
    carpeta_edificio = edificio["nombre"].replace(" ", "_")
    archivo_edificio = f"{carpeta_edificio}/{edificio['nombre']}_ascensores.json"

    if not os.path.exists(carpeta_edificio):
        os.makedirs(carpeta_edificio)

    if os.path.exists(archivo_edificio):
        with open(archivo_edificio, "r") as file:
            edificio = json.load(file)

    ascensor = {}
    ascensor["numero"] = int(entrada_numero_ascensor.get())
    ascensor["tipo_mantenimiento"] = combo_tipo_mantenimiento.get()
    if ascensor["tipo_mantenimiento"] == "Correctivo":
        correccion_seleccionada = combo_correccion.get()
        ascensor["correccion"] = correccion_seleccionada
        # Actualizar estadísticas
        if correccion_seleccionada in edificio["estadisticas_correccion"]:
            edificio["estadisticas_correccion"][correccion_seleccionada] += 1
        else:
            edificio["estadisticas_correccion"][correccion_seleccionada] = 1
    elif ascensor["tipo_mantenimiento"] == "Modernización":
        ascensor["valor"] = entrada_valor.get()
    ascensor["fecha_mantenimiento"] = cal.get_date()
    ascensor["descripcion_problema"] = entrada_descripcion_problema.get()

    edificio["ascensores"].append(ascensor)

    with open(archivo_edificio, "w") as file:
        json.dump(edificio, file, indent=4)

    generar_documento(edificio, carpeta_edificio)
    generar_estadisticas_excel(edificio, carpeta_edificio)

    messagebox.showinfo("Éxito", "Ascensor agregado correctamente.")

def generar_estadisticas_excel(edificio, carpeta_edificio):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Estadísticas de Corrección"
    ws.append(["Corrección", "Cantidad"])

    
    for cell in ws[1]:
        cell.font = Font(bold=True)

    
    ws.row_dimensions[1].height = 20

    
    ws.column_dimensions['A'].width = 30

    for correccion, cantidad in edificio["estadisticas_correccion"].items():
        ws.append([correccion, cantidad])

    estadisticas_path = os.path.join(carpeta_edificio, "estadisticas_correccion.xlsx")
    wb.save(estadisticas_path)

def seleccionar_tipo_mantenimiento(event):
    if combo_tipo_mantenimiento.get() == "Correctivo":
        etiqueta_correccion.grid(row=4, column=0, sticky="w")
        combo_correccion.grid(row=4, column=1, sticky="ew")
        etiqueta_valor.grid_forget()
        entrada_valor.grid_forget()
    elif combo_tipo_mantenimiento.get() == "Modernización":
        etiqueta_valor.grid(row=4, column=0, sticky="w")
        entrada_valor.grid(row=4, column=1, sticky="ew")
        etiqueta_correccion.grid_forget()
        combo_correccion.grid_forget()
    else:
        etiqueta_correccion.grid_forget()
        combo_correccion.grid_forget()
        etiqueta_valor.grid_forget()
        entrada_valor.grid_forget()

def mostrar_calendario(event):
    cal.grid(row=5, column=1, columnspan=2, sticky="ew")
    boton_guardar.grid(row=6, column=1, columnspan=2, sticky="ew")

def guardar_fecha():
    entrada_fecha_mantenimiento.delete(0, tk.END)
    entrada_fecha_mantenimiento.insert(0, cal.get_date())
    cal.grid_forget()
    boton_guardar.grid_forget()

def generar_documento(edificio, carpeta_edificio):
    document = Document()
    document.add_heading(f"Información del edificio {edificio['nombre']}", level=1)
    document.add_paragraph(f"NIT: {edificio['nit']}")
    document.add_heading("Ascensores", level=2)
    for ascensor in edificio["ascensores"]:
        document.add_paragraph(f"Número: {ascensor['numero']}")
        document.add_paragraph(f"Tipo de mantenimiento: {ascensor['tipo_mantenimiento']}")
        if ascensor["tipo_mantenimiento"] == "Correctivo":
            document.add_paragraph(f"Corrección: {ascensor['correccion']}")
        if ascensor["tipo_mantenimiento"] == "Modernización":
            valor = int(ascensor["valor"])
            if valor >= 1000000:
                valor_str = f"{valor} (Millones)"
            else:
                valor_str = f"{valor} (Miles)"
            document.add_paragraph(f"Valor: {valor_str}")
        document.add_paragraph(f"Fecha de mantenimiento: {ascensor['fecha_mantenimiento']}")
        document.add_paragraph(f"Descripción del problema: {ascensor['descripcion_problema']}")
        document.add_paragraph("--------------------------")
    documento_path = os.path.join(carpeta_edificio, "informacion_ascensores.docx")
    document.save(documento_path)
    messagebox.showinfo("Éxito", "Documento generado correctamente.")

ventana = tk.Tk()
ventana.title("Gestión de Ascensores")
ventana.resizable(True,True)
ventana.geometry("800x500")

datos = {}

edificios_nit, edificios_nombre = cargar_edificios_desde_excel("Listado de clientes.xlsx")

etiqueta_nombre_edificio = tk.Label(ventana, text="Nombre del edificio:")
etiqueta_nombre_edificio.grid(row=0, column=0, sticky="w")
combo_nombre_edificio = ttk.Combobox(ventana, values=edificios_nombre, state="readonly")
combo_nombre_edificio.grid(row=0, column=1, sticky="ew")
combo_nombre_edificio.bind("<<ComboboxSelected>>", actualizar_nit)

etiqueta_nit_edificio = tk.Label(ventana, text="NIT del edificio:")
etiqueta_nit_edificio.grid(row=1, column=0, sticky="w")
combo_nit_edificio = ttk.Combobox(ventana, values=list(edificios_nit.values()), state="readonly")
combo_nit_edificio.grid(row=1, column=1, sticky="ew")

etiqueta_numero_ascensor = tk.Label(ventana, text="Número de ascensor:")
etiqueta_numero_ascensor.grid(row=2, column=0, sticky="w")
opciones_ascensores = ["1", "2", "3", "4", "5"]
entrada_numero_ascensor = ttk.Combobox(ventana, values=opciones_ascensores, state="readonly")
entrada_numero_ascensor.grid(row=2, column=1, sticky="ew")

etiqueta_tipo_mantenimiento = tk.Label(ventana, text="Tipo de mantenimiento:")
etiqueta_tipo_mantenimiento.grid(row=3, column=0, sticky="w")
opciones_tipo_mantenimiento = ["Preventivo", "Correctivo", "Modernización"]
combo_tipo_mantenimiento = ttk.Combobox(ventana, values=opciones_tipo_mantenimiento, state="readonly")
combo_tipo_mantenimiento.grid(row=3, column=1, sticky="ew")
combo_tipo_mantenimiento.bind("<<ComboboxSelected>>", seleccionar_tipo_mantenimiento)

etiqueta_correccion = tk.Label(ventana, text="Corrección:")
opciones_correccion = [
    "Cambio de botón stop",
    "Falla en el fluido eléctrico",
    "Cambio de guaya",
    "Ajuste de puertas de cabina",
    "Cambio de indicador",
    "Bloqueo de tarjeta",
    "Ajuste gancho de puerta",
    "Cambio de contactor del freno",
    "Ajustes de cerraduras",
    "Cambio de lámparas",
    "Cambio de ventilador",
    "Instalación de fotocelda",
    "Cambio de tarjeta de llamada",
    "Otros"
]
combo_correccion = ttk.Combobox(ventana, values=opciones_correccion, state="readonly")

etiqueta_valor = tk.Label(ventana, text="Valor:")
entrada_valor = tk.Entry(ventana)

etiqueta_fecha_mantenimiento = tk.Label(ventana, text="Fecha de mantenimiento:")
etiqueta_fecha_mantenimiento.grid(row=5, column=0, sticky="w")
entrada_fecha_mantenimiento = tk.Entry(ventana)
entrada_fecha_mantenimiento.grid(row=5, column=1, sticky="ew")
entrada_fecha_mantenimiento.bind("<FocusIn>", mostrar_calendario)

etiqueta_descripcion_problema = tk.Label(ventana, text="Descripción del problema:")
etiqueta_descripcion_problema.grid(row=6, column=0, sticky="w")
entrada_descripcion_problema = tk.Entry(ventana)
entrada_descripcion_problema.grid(row=6, column=1, sticky="ew")

cal = Calendar(ventana, selectmode="day", year=2024, month=5, day=10)
cal.grid_forget()

boton_guardar = tk.Button(ventana, text="Guardar", command=guardar_fecha)
boton_guardar.grid_forget()

boton_agregar_ascensor = tk.Button(ventana, text="Agregar Ascensor", command=agregar_ascensor)
boton_agregar_ascensor.grid(row=7, column=0, columnspan=2, sticky="ew")

ventana.mainloop()